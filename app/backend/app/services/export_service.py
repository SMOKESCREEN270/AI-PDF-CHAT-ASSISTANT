"""
Export feature: turns any exportable payload (chat transcript, summary
dashboard, generated quiz/flashcards/questionnaire, comparison table) into
a downloadable file in the user's chosen format.
"""
import json
import os
from typing import Any, Dict
from datetime import datetime

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

EXPORT_DIR = "./exports"
os.makedirs(EXPORT_DIR, exist_ok=True)


def _payload_to_lines(kind: str, payload: Dict[str, Any]) -> list:
    """Flattens a payload dict into a simple list of (heading, body) lines for
    reuse across docx/pdf/markdown renderers."""
    lines = []
    if kind == "chat":
        for msg in payload.get("messages", []):
            speaker = "You" if msg["role"] == "user" else "Assistant"
            lines.append((speaker, msg["content"]))
            if msg.get("citations"):
                cite_str = "; ".join(
                    f"{c['filename']} p.{c['page']} (lines {c.get('line_start')}-{c.get('line_end')})"
                    for c in msg["citations"]
                )
                lines.append(("Sources", cite_str))
    elif kind in ("quiz", "flashcards"):
        items = payload.get("items", [])
        for i, item in enumerate(items, 1):
            if kind == "quiz":
                opts = "\n".join(f"  {chr(65+j)}. {o}" for j, o in enumerate(item.get("options", [])))
                lines.append((f"Q{i}. {item.get('question')}",
                               f"{opts}\nAnswer: {chr(65+item.get('correct_index', 0))} - {item.get('explanation','')}"))
            else:
                lines.append((f"Card {i}: {item.get('front')}", item.get('back', '')))
    elif kind == "questionnaire":
        for i, q in enumerate(payload.get("items", []), 1):
            lines.append((f"Q{i} [{q.get('type')}, {q.get('difficulty')}] {q.get('question')}",
                           f"Model answer: {q.get('model_answer','')} (source p.{q.get('source_page')})"))
    elif kind == "comparison":
        lines.append(("Dimensions", ", ".join(payload.get("dimensions", []))))
        for row in payload.get("table", []):
            values = "; ".join(f"{k}: {v}" for k, v in row.get("values", {}).items())
            lines.append((row.get("document", ""), values))
        for rec in payload.get("recommendations", []):
            lines.append((f"Best for: {rec.get('scenario')}", f"{rec.get('best_document')} - {rec.get('reason')}"))
    elif kind == "summary":
        lines.append(("Short Summary", payload.get("short_summary", "")))
        lines.append(("Key Insights", "\n".join(f"- {k}" for k in payload.get("key_insights", []))))
        lines.append(("Conclusion", payload.get("conclusion", "")))
    return lines


def export_to_json(kind: str, payload: Dict[str, Any], out_name: str) -> str:
    path = os.path.join(EXPORT_DIR, f"{out_name}.json")
    with open(path, "w") as f:
        json.dump(payload, f, indent=2)
    return path


def export_to_markdown(kind: str, payload: Dict[str, Any], out_name: str) -> str:
    lines = _payload_to_lines(kind, payload)
    path = os.path.join(EXPORT_DIR, f"{out_name}.md")
    with open(path, "w") as f:
        f.write(f"# {kind.title()} Export\n\n_Generated {datetime.utcnow().isoformat()}Z_\n\n")
        for heading, body in lines:
            f.write(f"## {heading}\n\n{body}\n\n")
    return path


def export_to_docx(kind: str, payload: Dict[str, Any], out_name: str) -> str:
    lines = _payload_to_lines(kind, payload)
    doc = DocxDocument()
    doc.add_heading(f"{kind.title()} Export", level=1)
    doc.add_paragraph(f"Generated {datetime.utcnow().isoformat()}Z")
    for heading, body in lines:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    path = os.path.join(EXPORT_DIR, f"{out_name}.docx")
    doc.save(path)
    return path


def export_to_pdf(kind: str, payload: Dict[str, Any], out_name: str) -> str:
    lines = _payload_to_lines(kind, payload)
    path = os.path.join(EXPORT_DIR, f"{out_name}.pdf")
    styles = getSampleStyleSheet()
    story = [Paragraph(f"{kind.title()} Export", styles["Title"]), Spacer(1, 0.2 * inch)]
    for heading, body in lines:
        story.append(Paragraph(heading, styles["Heading2"]))
        story.append(Paragraph(body.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 0.15 * inch))
    doc = SimpleDocTemplate(path, pagesize=LETTER)
    doc.build(story)
    return path


EXPORTERS = {
    "json": export_to_json,
    "markdown": export_to_markdown,
    "docx": export_to_docx,
    "pdf": export_to_pdf,
}


def export(kind: str, payload: Dict[str, Any], fmt: str, out_name: str) -> str:
    if fmt not in EXPORTERS:
        raise ValueError(f"Unsupported export format: {fmt}")
    return EXPORTERS[fmt](kind, payload, out_name)
